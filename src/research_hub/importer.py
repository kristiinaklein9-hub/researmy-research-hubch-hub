"""Local file import pipeline for non-DOI documents."""

from __future__ import annotations

import hashlib
import logging
import re
import sys
import warnings
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import yaml

from research_hub.clusters import ClusterRegistry, slugify
from research_hub.dedup import DedupHit, DedupIndex
from research_hub.manifest import Manifest, new_entry
from research_hub.security import atomic_write_text, safe_join, validate_slug

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS: dict[str, str] = {
    "pdf": "pdf",
    "md": "markdown",
    "markdown": "markdown",
    "txt": "txt",
    "docx": "docx",
    "url": "url",
}

_HASH_KEY_PREFIX = "hash:"
_BODY_PREVIEW_LIMIT = 5000
_SUMMARY_LIMIT = 500


@dataclass
class ImportEntry:
    path: Path
    slug: str = ""
    status: str = ""
    error: str = ""
    note_path: Path | None = None
    source_kind: str = ""


@dataclass
class ImportReport:
    folder: Path
    cluster_slug: str
    entries: list[ImportEntry] = field(default_factory=list)
    dry_run: bool = False

    @property
    def imported_count(self) -> int:
        return sum(1 for entry in self.entries if entry.status == "imported")

    @property
    def skipped_count(self) -> int:
        return sum(1 for entry in self.entries if entry.status.startswith("skipped"))

    @property
    def failed_count(self) -> int:
        return sum(1 for entry in self.entries if entry.status == "failed")


# v0.88.14: content-hash cache for PDF text extraction.
#
# pdfplumber walking every page of a non-trivial PDF (~20 pages) takes
# 1-3 s and produces ~50 KB of plain text. Re-importing the same PDF
# (which happens whenever the user retries an ingest, re-runs auto on
# the same cluster, or moves a paper between clusters) re-paid that
# cost every time, with no caching layer anywhere in the pipeline.
#
# We key by sha256 of the file bytes so a moved/renamed PDF (same
# content) is still a cache hit. The cache lives under
# `<vault>/.research_hub/cache/pdf_extract/<sha>.txt` — vault-scoped
# so it's GC'able alongside the rest of the sidecars. Set via
# `set_pdf_extract_cache_dir(path)` at pipeline boot; if unset, no
# caching happens (default-safe for tests + ad-hoc importer use).

_PDF_EXTRACT_CACHE_DIR: Path | None = None


def set_pdf_extract_cache_dir(path: Path | None) -> None:
    """v0.88.14: pipeline init point — tell ``_extract_pdf`` where to
    persist its content-hash cache. Pass ``None`` to disable caching."""
    global _PDF_EXTRACT_CACHE_DIR
    _PDF_EXTRACT_CACHE_DIR = Path(path) if path else None


def _pdf_cache_paths(file_path: Path) -> tuple[Path | None, str]:
    """Return (cache_path, sha256_hex) — cache_path is None when
    caching is disabled."""
    import hashlib

    cache_dir = _PDF_EXTRACT_CACHE_DIR
    digest = hashlib.sha256(file_path.read_bytes()).hexdigest()
    if cache_dir is None:
        return None, digest
    return cache_dir / f"{digest}.txt", digest


def _extract_pdf(path: Path) -> str:
    # v0.88.14: hash + cache hit short-circuit before invoking pdfplumber
    cache_path, _digest = _pdf_cache_paths(path)
    if cache_path is not None and cache_path.exists():
        try:
            return cache_path.read_text(encoding="utf-8")
        except OSError:
            pass  # fall through to fresh extraction

    try:
        import pdfplumber
    except ImportError as exc:  # pragma: no cover - dependency-gated
        raise RuntimeError(
            "PDF extraction requires pdfplumber. Install: pip install 'research-hub-pipeline[import]'"
        ) from exc

    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text)
    extracted = "\n\n".join(parts).strip()

    # v0.88.14: write cache (best-effort; never let a cache failure
    # poison the extraction itself)
    if cache_path is not None:
        try:
            cache_path.parent.mkdir(parents=True, exist_ok=True)
            cache_path.write_text(extracted, encoding="utf-8")
        except OSError:
            pass

    return extracted


def _pdf_metadata_title(path: Path) -> str | None:
    try:
        import pdfplumber
    except ImportError as exc:  # pragma: no cover - dependency-gated
        raise RuntimeError(
            "PDF extraction requires pdfplumber. Install: pip install 'research-hub-pipeline[import]'"
        ) from exc

    with pdfplumber.open(str(path)) as pdf:
        metadata = pdf.metadata or {}
        title = str(metadata.get("Title", "") or "").strip()
    return title or None


def _extract_markdown(path: Path) -> str:
    text = path.read_text(encoding="utf-8", errors="replace")
    if text.startswith("---\n"):
        end = text.find("\n---\n", 4)
        if end > 0:
            text = text[end + 5 :]
    return text.strip()


def _extract_docx(path: Path) -> tuple[str, str]:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dependency-gated
        raise RuntimeError(
            "DOCX extraction requires python-docx. Install: pip install 'research-hub-pipeline[import]'"
        ) from exc

    doc = docx.Document(str(path))
    title = str(doc.core_properties.title or "").strip()
    if not title:
        for paragraph in doc.paragraphs[:5]:
            style_name = getattr(getattr(paragraph, "style", None), "name", "")
            if style_name in {"Heading 1", "Title"} and paragraph.text.strip():
                title = paragraph.text.strip()
                break
    body = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()
    return title, body


def _extract_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="replace").strip()
    except OSError as exc:
        raise RuntimeError(f"could not read {path}: {exc}") from exc


def _extract_url(path: Path) -> str:
    first_line = next((line.strip() for line in path.read_text(encoding="utf-8").splitlines() if line.strip()), "")
    if not first_line.startswith(("http://", "https://")):
        raise ValueError(f"{path}: first non-empty line must be a URL")

    # Prefer defuddle (v0.43): cleaner output, actively maintained.
    from research_hub.defuddle_extract import extract_url_via_defuddle

    defuddle_result = extract_url_via_defuddle(first_line)
    if defuddle_result is not None:
        return defuddle_result

    # Fallback: readability-lxml (v0.42 default; unmaintained since 2021).
    try:
        import requests
        from readability import Document as ReadabilityDocument
    except ImportError as exc:  # pragma: no cover - dependency-gated
        raise RuntimeError(
            "URL extraction requires either defuddle CLI or "
            "readability-lxml. Install: npm install -g defuddle-cli, OR "
            "pip install 'research-hub-pipeline[import]'"
        ) from exc

    response = requests.get(
        first_line,
        timeout=30,
        headers={"User-Agent": "research-hub/0.43"},
    )
    response.raise_for_status()
    html_summary = ReadabilityDocument(response.text).summary()
    return _html_to_text(html_summary)


def _html_to_text(html: str) -> str:
    """Strip HTML tags from readability output, preserving paragraph breaks.

    readability-lxml returns the article body as HTML — we want plain text in
    the Obsidian note. Uses html.parser (stdlib only).
    """
    from html.parser import HTMLParser

    class _TextCollector(HTMLParser):
        BLOCK_TAGS = {"p", "div", "h1", "h2", "h3", "h4", "h5", "h6", "br", "li"}

        def __init__(self) -> None:
            super().__init__()
            self.parts: list[str] = []

        def handle_starttag(self, tag: str, attrs):
            if tag in self.BLOCK_TAGS and self.parts and not self.parts[-1].endswith("\n"):
                self.parts.append("\n")

        def handle_endtag(self, tag: str):
            if tag in self.BLOCK_TAGS:
                self.parts.append("\n")

        def handle_data(self, data: str):
            if data.strip():
                self.parts.append(data.strip())

    collector = _TextCollector()
    try:
        collector.feed(html)
    except Exception:
        return html  # fallback to raw if parser blows up
    text = " ".join(collector.parts)
    # Collapse runs of whitespace + collapse repeated newlines
    import re
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n[ \t]*\n+", "\n\n", text)
    return text.strip()


_EXTRACTORS = {
    "pdf": _extract_pdf,
    "markdown": _extract_markdown,
    "txt": _extract_txt,
    "url": _extract_url,
}


def _first_nonempty_line(text: str) -> str:
    return next((line.strip() for line in text.splitlines() if line.strip()), "")


def _derive_title(text: str, fallback: Path, source_kind: str, source_path: Path) -> str:
    if source_kind == "markdown":
        for line in text.splitlines()[:50]:
            stripped = line.strip()
            if stripped.startswith("# "):
                return stripped[2:].strip()
    elif source_kind == "pdf":
        metadata_title = _pdf_metadata_title(source_path)
        if metadata_title:
            return metadata_title
        first_line = _first_nonempty_line(text)
        if first_line:
            return first_line
    elif source_kind == "txt":
        first_line = _first_nonempty_line(text)
        if first_line and len(first_line) < 100:
            return first_line
    return fallback.stem.replace("_", " ").replace("-", " ").strip() or fallback.stem


def _derive_title_for_kind(text: str, source_path: Path, source_kind: str, *, title_hint: str = "") -> str:
    if title_hint.strip():
        return title_hint.strip()
    return _derive_title(text, source_path, source_kind, source_path)


def _filename_slug(path: Path) -> str:
    base = re.sub(r"[^a-z0-9]+", "-", path.stem.lower()).strip("-")
    return base[:64] or "imported-file"


def _derive_slug(title: str, source_path: Path) -> str:
    candidate = slugify(title)
    candidate = re.sub(r"[^a-z0-9_-]+", "-", candidate.lower()).strip("-_")
    return (candidate[:64] if candidate else _filename_slug(source_path)) or "imported-file"


def _content_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8", errors="replace")).hexdigest()[:16]


def _truncate_body(text: str) -> str:
    if len(text) <= _BODY_PREVIEW_LIMIT:
        return text
    return text[:_BODY_PREVIEW_LIMIT] + "\n\n*(truncated; see raw_path for the source file)*"


def _document_frontmatter(
    *,
    slug_value: str,
    title: str,
    source_kind: str,
    cluster_slug: str,
    raw_path: Path,
    summary: str,
) -> dict[str, Any]:
    try:
        from research_hub.document import Document

        doc = Document(
            slug=slug_value,
            title=title,
            source_kind=source_kind,
            topic_cluster=cluster_slug,
            ingestion_source="import-folder",
            raw_path=str(raw_path),
            summary=summary,
        )
        if hasattr(doc, "to_frontmatter"):
            frontmatter = dict(doc.to_frontmatter())
        else:
            frontmatter = {
                "title": title,
                "slug": slug_value,
                "source_kind": source_kind,
                "topic_cluster": cluster_slug,
                "ingestion_source": "import-folder",
                "raw_path": str(raw_path),
                "summary": summary,
            }
    except Exception:
        frontmatter = {
            "title": title,
            "slug": slug_value,
            "source_kind": source_kind,
            "topic_cluster": cluster_slug,
            "ingested_at": "",
            "ingestion_source": "import-folder",
            "labels": [],
            "tags": [],
            "raw_path": str(raw_path),
            "summary": summary,
        }

    frontmatter.setdefault("slug", slug_value)
    frontmatter["title"] = title
    frontmatter["source_kind"] = source_kind
    frontmatter["topic_cluster"] = cluster_slug
    frontmatter["ingestion_source"] = "import-folder"
    frontmatter["raw_path"] = str(raw_path)
    frontmatter["summary"] = summary
    frontmatter.setdefault("labels", [])
    frontmatter.setdefault("tags", [])
    frontmatter.setdefault("cluster_queries", [])
    frontmatter.setdefault("status", "unread")
    return frontmatter


def _render_document_md(
    *,
    slug_value: str,
    title: str,
    source_kind: str,
    cluster_slug: str,
    raw_path: Path,
    text: str,
) -> str:
    frontmatter = _document_frontmatter(
        slug_value=slug_value,
        title=title,
        source_kind=source_kind,
        cluster_slug=cluster_slug,
        raw_path=raw_path,
        summary=text[:_SUMMARY_LIMIT],
    )
    yaml_str = yaml.safe_dump(frontmatter, allow_unicode=True, sort_keys=False)
    body = _truncate_body(text).strip()
    return f"---\n{yaml_str}---\n\n{body}\n"


def _maybe_create_cluster(cfg: Any, cluster_slug: str, *, dry_run: bool) -> None:
    registry = ClusterRegistry(cfg.clusters_file)
    if registry.get(cluster_slug) is not None or dry_run:
        return
    registry.create(
        query=cluster_slug,
        name=cluster_slug.replace("-", " ").title(),
        slug=cluster_slug,
    )
    logger.info("auto-created cluster: %s", cluster_slug)


def _load_dedup(path: Path) -> DedupIndex:
    return DedupIndex.load(path) if path.exists() else DedupIndex()


def _hash_key(content_hash: str) -> str:
    return f"{_HASH_KEY_PREFIX}{content_hash}"


def import_folder(
    cfg: Any,
    folder: Path | str,
    *,
    cluster_slug: str,
    extensions: tuple[str, ...] | None = None,
    skip_existing: bool = True,
    use_graphify: bool = False,
    graphify_graph: Path | None = None,
    dry_run: bool = False,
    with_zotero: bool = False,
    yes: bool = False,
    batch_label: str | None = None,
) -> ImportReport:
    cluster_slug = validate_slug(cluster_slug, field="cluster_slug")
    if not with_zotero and not dry_run:
        msg = (
            "WARNING: import-folder writes Obsidian only.\n"
            f"  Cluster '{cluster_slug}' Zotero collection will NOT receive these files.\n"
            "  To also write Zotero, re-run with --with-zotero (requires Zotero API set up).\n"
        )
        print(msg, file=sys.stderr)
        if not yes and getattr(sys.stdin, "isatty", lambda: False)():
            try:
                response = input("Continue with Obsidian-only? [y/N]: ").strip().lower()
            except EOFError:
                response = ""
            if response not in {"y", "yes"}:
                raise SystemExit("aborted by user")
    folder_path = Path(folder).expanduser().resolve()
    if not folder_path.is_dir():
        raise ValueError(f"folder not found: {folder_path}")

    requested_extensions = extensions or ("pdf", "md", "txt", "docx", "url")
    normalized_extensions = tuple(ext.lower().lstrip(".") for ext in requested_extensions)

    _maybe_create_cluster(cfg, cluster_slug, dry_run=dry_run)
    # v0.88.14: opt the PDF extract into vault-scoped content-hash
    # caching for this run. Re-imports of the same PDF (rename, retry,
    # cluster move) skip pdfplumber entirely.
    set_pdf_extract_cache_dir(cfg.research_hub_dir / "cache" / "pdf_extract")
    dedup_path = cfg.research_hub_dir / "dedup_index.json"
    dedup = _load_dedup(dedup_path)
    manifest = Manifest(cfg.research_hub_dir / "manifest.jsonl")
    cluster_raw_dir = safe_join(cfg.raw, cluster_slug)
    if not dry_run:
        cluster_raw_dir.mkdir(parents=True, exist_ok=True)

    report = ImportReport(folder=folder_path, cluster_slug=cluster_slug, dry_run=dry_run)
    imported_records: list[dict[str, Any]] = []

    if use_graphify and not graphify_graph:
        warnings.warn(
            "--use-graphify is deprecated in v0.32. graphify cannot be invoked "
            "as a standalone CLI. Run /graphify in Claude Code to produce "
            "graph.json, then pass --graphify-graph PATH instead. "
            "See docs/import-folder.md.",
            DeprecationWarning,
            stacklevel=2,
        )

    for path in sorted(folder_path.rglob("*")):
        if not path.is_file():
            continue

        ext = path.suffix.lower().lstrip(".")
        if ext not in normalized_extensions or ext not in SUPPORTED_EXTENSIONS:
            report.entries.append(ImportEntry(path=path, status="skipped_unsupported"))
            continue

        source_kind = SUPPORTED_EXTENSIONS[ext]
        entry = ImportEntry(path=path, source_kind=source_kind)
        try:
            title_hint = ""
            if source_kind == "docx":
                title_hint, text = _extract_docx(path)
            else:
                text = _EXTRACTORS[source_kind](path)
        except Exception as exc:
            entry.status = "failed"
            entry.error = f"extraction failed: {exc}"
            report.entries.append(entry)
            continue

        title = _derive_title_for_kind(text, path, source_kind, title_hint=title_hint)
        entry.slug = _derive_slug(title, path)
        content_hash = _content_hash(text)
        if skip_existing and dedup.title_to_hits.get(_hash_key(content_hash)):
            entry.status = "skipped_duplicate"
            report.entries.append(entry)
            continue

        markdown = _render_document_md(
            slug_value=entry.slug,
            title=title,
            source_kind=source_kind,
            cluster_slug=cluster_slug,
            raw_path=path,
            text=text,
        )

        if not dry_run:
            note_path = cluster_raw_dir / f"{entry.slug}.md"
            if note_path.exists():
                note_path = cluster_raw_dir / f"{entry.slug}-{content_hash[:8]}.md"
            atomic_write_text(note_path, markdown)
            entry.note_path = note_path
            dedup.title_to_hits[_hash_key(content_hash)] = [
                DedupHit(
                    source="importer",
                    title=title,
                    obsidian_path=str(note_path),
                )
            ]
            imported_records.append(
                {
                    "entry": entry,
                    "title": title,
                    "content_hash": content_hash,
                    "note_path": note_path,
                    "text": text,
                }
            )

        entry.status = "imported"
        report.entries.append(entry)

    if not dry_run and report.imported_count > 0:
        dedup.save(dedup_path)
        manifest_query = batch_label or "import-folder"
        manifest_batch_label = batch_label or ""
        zotero_keys = [""] * len(imported_records)
        if with_zotero and imported_records:
            from research_hub.pipeline import (
                _ensure_batch_subcollection,
                resolve_batch_label,
                write_papers_to_zotero,
            )
            from research_hub.zotero.client import get_client

            zot = get_client()
            cluster = ClusterRegistry(cfg.clusters_file).get(cluster_slug)
            cluster_coll = cluster.zotero_collection_key if cluster else None
            zotero_batch_label = resolve_batch_label(None, batch_label)
            batch_coll = ""
            if cluster_coll:
                batch_coll = _ensure_batch_subcollection(
                    zot,
                    cluster_coll=cluster_coll,
                    batch_label=zotero_batch_label,
                    log=logger.info,
                )
            import_papers = [
                {
                    "title": record["title"],
                    "doi": _hash_key(record["content_hash"]),
                    "authors": [],
                    "year": datetime.now(timezone.utc).year,
                    "abstract": "",
                    "journal": "(local file)",
                    "slug": record["entry"].slug,
                    "sub_category": cluster_slug,
                    "summary": str(record["text"])[:500],
                    "key_findings": [],
                    "methodology": "",
                    "relevance": "",
                    "url": "",
                }
                for record in imported_records
            ]
            _zr, _papers_for_notes, errors = write_papers_to_zotero(
                zot,
                import_papers,
                cluster_slug,
                cluster_coll,
                batch_coll=batch_coll or None,
                batch_label=zotero_batch_label if cluster_coll else "",
                log=logger.info,
            )
            if errors:
                logger.warning("import-folder Zotero write had %d error(s)", len(errors))
            zotero_keys = [paper.get("zotero_key", "") for paper in import_papers]
            manifest_batch_label = zotero_batch_label
        for record, zotero_key in zip(imported_records, zotero_keys):
            manifest.append(
                new_entry(
                    cluster=cluster_slug,
                    query=manifest_query,
                    action="import-folder-with-zotero" if with_zotero else "import-folder",
                    title=record["title"],
                    zotero_key=zotero_key,
                    obsidian_path=str(record["note_path"]),
                    batch_label=manifest_batch_label,
                )
            )

    if graphify_graph and not dry_run:
        _apply_graphify_assignments(report, Path(graphify_graph))

    return report


def _apply_graphify_assignments(report: ImportReport, graphify_graph: Path) -> None:
    try:
        from research_hub.graphify_bridge import (
            map_to_subtopics,
            parse_graphify_communities,
        )
    except ImportError:
        logger.warning("graphify_bridge not available; skipping graphify graph import")
        return

    try:
        communities = parse_graphify_communities(graphify_graph)
        assignments = map_to_subtopics(
            communities,
            [entry.path for entry in report.entries if entry.status == "imported"],
        )
        for entry in report.entries:
            if entry.status != "imported" or entry.note_path is None:
                continue
            subtopics = assignments.get(str(entry.path), [])
            if subtopics:
                _add_subtopics_frontmatter(entry.note_path, subtopics)
        logger.info(
            "applied %d community assignments from %s",
            sum(1 for value in assignments.values() if value),
            graphify_graph,
        )
    except Exception as exc:  # pragma: no cover - best effort bridge
        logger.warning("graphify graph.json parse failed: %s", exc)


def _add_subtopics_frontmatter(note_path: Path, subtopics: list[str]) -> None:
    text = note_path.read_text(encoding="utf-8")
    if not text.startswith("---\n"):
        return
    end = text.find("\n---\n", 4)
    if end < 0:
        return
    frontmatter = yaml.safe_load(text[4:end]) or {}
    if not isinstance(frontmatter, dict):
        return
    frontmatter["subtopics"] = list(subtopics)
    yaml_str = yaml.safe_dump(frontmatter, allow_unicode=True, sort_keys=False)
    body = text[end + 5 :]
    note_path.write_text(f"---\n{yaml_str}---\n{body}", encoding="utf-8")


__all__ = ["import_folder", "ImportEntry", "ImportReport", "SUPPORTED_EXTENSIONS"]
